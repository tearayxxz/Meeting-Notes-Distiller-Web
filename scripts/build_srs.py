from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Meeting_Notes_Distiller_SRS.docx"
BLUE = "2E74B5"
DARK = "17324D"
LIGHT = "F2F4F7"
BORDER = "C7D1DB"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    props = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    props.append(repeat)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 24, DARK, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, DARK, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    list_style = document.styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_after = Pt(8)
    list_style.paragraph_format.line_spacing = 1.167


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "MEETING NOTES DISTILLER WEB  |  SOFTWARE REQUIREMENTS SPECIFICATION"
    header.style = "Caption"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.color.rgb = RGBColor.from_string("64748B")
    header.runs[0].font.size = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("University assignment deliverable  •  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("64748B")
    add_page_number(footer)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK)
        if widths:
            cell.width = Inches(widths[index])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[index])
    document.add_paragraph()
    return table


def add_requirement_table(document, rows):
    return add_table(document, ["ID", "Requirement", "Implemented behavior"], rows, [0.75, 2.65, 3.6])


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(document)
    add_header_footer(section)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Meeting Notes Distiller Web")
    subtitle = document.add_paragraph()
    subtitle.add_run("Software Requirements Specification").bold = True
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    document.add_paragraph("Version 1.0  |  Implementation baseline  |  24 August 2026")
    document.add_paragraph(
        "A deterministic, offline web application for converting plain-text meeting transcripts into structured, evidence-bound meeting intelligence and downloadable Word reports."
    )
    add_table(
        document,
        ["Document control", "Value"],
        [
            ("Product", "Meeting Notes Distiller Web"),
            ("Implementation", "React/Vite frontend, Express backend, shared TypeScript/Zod contracts"),
            ("Processing mode", "Offline deterministic rules and heuristics; no API keys"),
            ("Verification", "Vitest, Supertest, Testing Library, Playwright, OOXML checks"),
        ],
        [2.0, 5.0],
    )

    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(
        "This SRS defines the implemented behavior and acceptance boundary of Meeting Notes Distiller Web. The system accepts meeting transcript text files, normalizes several common speaker layouts, extracts structured information, highlights unresolved work, and produces a combined Microsoft Word report."
    )

    document.add_heading("2. Purpose", level=1)
    document.add_paragraph(
        "The product reduces the manual effort required to review meeting transcripts while keeping every significant result traceable to source text. It is designed for reproducible university assessment and local use rather than probabilistic cloud summarization."
    )

    document.add_heading("3. Scope", level=1)
    add_bullets(document, [
        "Accept up to 10 plain-text transcripts per request, each up to 1 MiB.",
        "Normalize speaker-colon, timestamp-dash, timestamp-block, and unstructured text.",
        "Extract participants, topics, topic summaries, explicit decisions, action items, owners, deadlines, and evidence.",
        "Flag unresolved options, conflicting dates, unassigned work, empty content, and parser warnings.",
        "Display per-meeting and combined results and generate a valid .docx report.",
        "Exclude authentication, database persistence, cloud infrastructure, paid APIs, and automatic calendar integration.",
    ])

    document.add_heading("4. Product Overview", level=1)
    document.add_paragraph(
        "The browser presents a responsive analysis dashboard. Users build a local queue across multiple selection rounds, explicitly start analysis, review successful and failed files, and export the current batch. The server holds files only in memory and returns typed JSON or a generated Word document."
    )
    add_table(document, ["Layer", "Responsibility"], [
        ("Frontend", "File queue, drag/drop, validation feedback, API calls, result cards, grouped actions, report download."),
        ("Backend", "Multipart limits, per-file isolation, normalization, extraction orchestration, safe errors, Word generation."),
        ("Shared", "Zod schemas and TypeScript types for normalized transcripts, meetings, failures, and report requests."),
        ("Tests", "Domain, schema, API, component, artifact, and browser-level observable behavior."),
    ], [1.2, 5.8])

    document.add_heading("5. User Roles / Actors", level=1)
    add_table(document, ["Actor", "Goal", "Permissions"], [
        ("Meeting reviewer", "Turn transcripts into a concise, auditable work summary.", "Select local files, analyze, review, and download."),
        ("Instructor/evaluator", "Verify required behavior and reproducibility.", "Run samples, tests, build, E2E, and inspect artifacts."),
        ("Local application", "Process files without external services.", "Reads request memory, returns results, stores nothing persistently."),
    ], [1.4, 3.1, 2.5])

    document.add_heading("6. Functional Requirements", level=1)
    add_requirement_table(document, [
        ("FR-01", "Accept one or more .txt files and allow later additions.", "Custom drag/drop and file input maintain an additive, deduplicated queue."),
        ("FR-02", "Require an explicit processing action.", "Analyze Meetings remains disabled until valid files exist and starts the request on click."),
        ("FR-03", "Reject unsupported, empty, invalid, oversized, or excessive uploads safely.", "A 10 MiB transport cap protects memory; per-file 1 MiB, UTF-8, empty, and type checks preserve valid peers."),
        ("FR-04", "Support at least three transcript layouts.", "Speaker-colon, timestamp-dash, and timestamp-block normalization plus unstructured fallback."),
        ("FR-05", "Extract unique participants from speaker information only.", "Participant module deduplicates normalized non-null speakers."),
        ("FR-06", "Identify topics and separate concise summaries.", "Fixed evidence-bound subject categories group matching utterances; unmatched text is General discussion."),
        ("FR-07", "Identify explicit final decisions and avoid promoting suggestions.", "Resolution phrases produce decisions with evidence; option language alone does not."),
        ("FR-08", "Extract action task, owner, deadline, and evidence without invention.", "Pattern modules return nullable owner/date; UI displays Unassigned and Not specified."),
        ("FR-09", "Flag meetings with unresolved options or conflicting dates.", "No-decision and date-conflict flags require multiple alternatives and no explicit decision."),
        ("FR-10", "Show clearly separate results for every successful meeting.", "Tabbed dashboard renders filename, participants, topics, decisions, actions, and warnings."),
        ("FR-11", "Group all actions by responsible person.", "Backend returns stable owner groups with Unassigned last; frontend renders owner cards."),
        ("FR-12", "Generate a valid combined Word report.", "POST /api/report validates the batch and returns OOXML .docx with all required sections."),
    ])

    document.add_heading("7. Non-Functional Requirements", level=1)
    add_requirement_table(document, [
        ("NFR-01", "Reproducible offline operation", "Core analysis requires no network, LLM, account, database, or secret."),
        ("NFR-02", "Type safety", "Strict TypeScript plus Zod validation at shared API/report boundaries."),
        ("NFR-03", "Usability and accessibility", "Responsive layout, keyboard focus, semantic controls, labels, high-contrast warnings, and loading state."),
        ("NFR-04", "Fault containment", "A malformed or empty file cannot abort analysis of other valid files in the batch."),
        ("NFR-05", "Security", "Memory-only upload handling, extension/size/count limits, no stack traces, and no secrets."),
        ("NFR-06", "Maintainability", "Normalization, extraction concerns, routes, UI, and contracts remain separate and testable."),
    ])

    document.add_heading("8. User Flow", level=1)
    add_bullets(document, [
        "Open the dashboard and review the empty upload state.",
        "Drag or select one or more .txt transcripts; optionally add more in later rounds.",
        "Review the queue and remove unwanted entries.",
        "Select Analyze Meetings and wait for the processing indicator.",
        "Review per-meeting results, global actions, warnings, and any failed filenames.",
        "Select Download Word Report to save the current successful analysis batch.",
    ])

    document.add_heading("9. System Architecture", level=1)
    document.add_paragraph(
        "The repository is an npm workspace with explicit frontend, backend, shared, tests, and E2E boundaries. Vite proxies development API calls to Express. A production build lets Express serve the generated frontend while retaining the same /api endpoints."
    )
    add_table(document, ["Component", "Primary implementation"], [
        ("Dashboard", "frontend/src/App.tsx and focused presentation components"),
        ("Contracts", "shared/src/contracts.ts"),
        ("Normalization", "backend/src/domain/normalizer.ts"),
        ("Semantics", "participants.ts, topics.ts, decisions.ts, actions.ts, flags.ts"),
        ("Batch service", "backend/src/services/analyze-files.ts"),
        ("Runtime report", "backend/src/services/report.ts"),
        ("HTTP boundary", "backend/src/routes/analyze.ts and report.ts"),
    ], [2.0, 5.0])

    document.add_heading("10. Processing / Data Flow", level=1)
    flow = document.add_paragraph()
    run = flow.add_run(
        "Uploaded TXT → input validation → UTF-8 text normalization → format detection → normalized utterances → participants/topics/decisions/actions/flags → shared AnalysisBatch → dashboard or Word report"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    document.add_paragraph(
        "Normalization retains line numbers, timestamps, nullable speakers, and utterance text. Semantic modules consume only the normalized representation. Deterministic meeting IDs hash the filename and normalized content. The batch response includes successful meetings, named failures, grouped actions, and an ISO processing timestamp."
    )

    document.add_heading("11. Error Handling", level=1)
    add_table(document, ["Condition", "System response"], [
        ("No files", "HTTP 400 and a safe request to select at least one .txt transcript."),
        ("Unsupported extension", "File-specific unsupported-type failure; client also rejects it before upload."),
        ("Empty text", "File-specific empty-file failure; other valid meetings remain."),
        ("Unknown format", "Unstructured normalization preserves text with null speaker and a parse warning when appropriate."),
        ("Malformed controls", "Text cleanup and guarded per-file processing avoid a batch crash."),
        ("Report payload invalid", "HTTP 400 with a safe validation message and no stack trace."),
        ("Unexpected server error", "Central middleware returns a generic response without exposing a stack trace."),
        ("Download/network failure", "Dashboard keeps analyzed results and displays a visible error."),
    ], [2.1, 4.9])

    document.add_heading("12. Testing Requirements", level=1)
    add_bullets(document, [
        "Unit tests cover all three required formats, unstructured and empty input, duplicate participants, explicit/non-decisions, multiple actions, missing owners, Thai patterns, and conflicts.",
        "Shared-schema tests reject malformed structured output.",
        "Supertest verifies no-file errors, successful analysis, partial failure, empty-file isolation, health, invalid report requests, and readable DOCX output.",
        "Testing Library verifies empty state, unsupported files, additive upload, duplicate suppression, and explicit analysis.",
        "Playwright verifies successful analysis, multiple upload rounds, grouped ownership, no-decision warning, unsupported extension, and report download.",
        "Artifact tests and a standalone verifier inspect required OOXML entries and representative document text.",
    ])

    document.add_heading("13. Assumptions", level=1)
    add_bullets(document, [
        "Input is valid UTF-8 plain text; invalid byte sequences are reported as a file-specific upload failure.",
        "Speaker names appear in supported line headers when participant extraction is expected.",
        "Decision language explicitly signals agreement or resolution.",
        "Relative deadlines such as Friday are meaningful to the user and should remain as written.",
        "The local browser can download Blob responses and open Office Open XML files with a compatible application.",
    ])

    document.add_heading("14. Constraints", level=1)
    add_bullets(document, [
        "At most 10 files are accepted. The transport cap is 10 MiB per file and the application transcript limit is 1 MiB.",
        "No persistent storage: refreshing the page clears the queue and results.",
        "The fixed topic vocabulary and phrase rules prioritize transparency over broad language understanding.",
        "The application does not resolve relative dates against a meeting date.",
        "The application does not call external services or use API keys for extraction.",
    ])

    document.add_heading("15. Acceptance Criteria", level=1)
    add_requirement_table(document, [
        ("AC-01", "One, multiple, and additive .txt selections can be explicitly analyzed.", "Covered by component, API, and Playwright tests."),
        ("AC-02", "Formats A, B, and C normalize without crashing.", "Covered by focused normalizer tests and sample data."),
        ("AC-03", "Required semantic sections render per meeting and globally.", "Covered by extractor tests and successful browser flow."),
        ("AC-04", "Ownerless requirement becomes Unassigned with Friday retained.", "Covered by mandatory action test and grouped browser flow."),
        ("AC-05", "Three launch options without resolution produce no decision warning.", "Covered by mandatory flag test and browser flow."),
        ("AC-06", "One failed file does not remove successful results.", "Covered by API partial-success tests."),
        ("AC-07", "Downloaded Word report is structurally valid and contains actual output.", "Covered by API, Playwright, and OOXML integrity checks."),
        ("AC-08", "Install, lint, tests, build, and E2E execute with documented commands.", "Verified during submission-readiness review."),
    ])

    document.add_heading("16. Known Limitations", level=1)
    add_bullets(document, [
        "Unusual phrasing can be missed because the engine is heuristic rather than a general language model.",
        "English has the broadest coverage; Thai support is limited to tested owner, weekday, and decision patterns.",
        "Conflict detection targets competing dates around shared release/launch concepts, not arbitrary contradictions.",
        "Topic names come from a fixed vocabulary; unmatched content is grouped as General discussion.",
        "Topic summaries are short evidence excerpts and may be less fluent than generated prose.",
        "The system does not perform diarization, OCR, audio transcription, user management, collaboration, or persistence.",
    ])

    document.core_properties.title = "Meeting Notes Distiller Web — Software Requirements Specification"
    document.core_properties.subject = "Implemented product and software requirements"
    document.core_properties.author = "Meeting Notes Distiller Web Project"
    document.core_properties.keywords = "SRS, meeting transcript, extraction, testing"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
